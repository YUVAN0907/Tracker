from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_document():
    doc = Document()
    
    # Title
    title = doc.add_heading('VendBees Application Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Interactive Components, Calculations & Improvements Guide')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Date: February 16, 2026')
    doc.add_paragraph()
    
    # ===================== SECTION 1: OVERVIEW =====================
    doc.add_heading('1. Application Overview', level=1)
    doc.add_paragraph(
        'VendBees is a vending machine inventory management system that tracks stock levels, '
        'sales, purchases, and machine performance in real-time. The application syncs data '
        'from a SharePoint Excel file and provides a web-based dashboard for monitoring.'
    )
    
    # Architecture
    doc.add_heading('Data Flow Architecture', level=2)
    arch = doc.add_paragraph()
    arch.add_run('SharePoint Excel').bold = True
    arch.add_run(' → (downloads every 30s) → ')
    arch.add_run('Backend Server (Flask @ Port 3001)').bold = True
    arch.add_run(' → ')
    arch.add_run('Frontend (React @ Port 5173/5174)').bold = True
    arch.add_run(' → (fetches every 5s) → ')
    arch.add_run('User Dashboard').bold = True
    
    # ===================== SECTION 2: DASHBOARD KPIs =====================
    doc.add_heading('2. Dashboard Page - KPI Cards', level=1)
    
    # KPI Table
    table = doc.add_table(rows=7, cols=4)
    table.style = 'Table Grid'
    
    # Header row
    headers = ['KPI Component', 'Calculation Formula', 'Real-World Business Logic', 'Interactive Filter']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')
    
    # Data rows
    kpi_data = [
        ['Total Stock Value', 'Σ (Current_Stock × PO_Cost) for all products', 
         'Shows how much capital is tied up in unsold inventory. Critical for cash flow management and investment decisions.',
         'Machine Dropdown - Filter by specific machine'],
        ['Total Units', 'Σ Current_Stock across all stock entries',
         'Total physical items available for sale. Helps plan logistics, storage, and restocking schedules.',
         'Machine Dropdown'],
        ['Out of Stock', 'Count where Current_Stock <= 0',
         'Products unavailable for sale = direct revenue loss. These need urgent restocking priority.',
         'Machine Dropdown'],
        ["Today's Sales", 'Σ (Qty × Selling_Price) where Date = today',
         'Daily revenue tracking for performance monitoring. Compare against targets and historical averages.',
         'Machine Dropdown'],
        ['Total GST Payable', 'Σ (Sale_Amount × GST_Rate) for all sales',
         'Tax liability to government under GST law. Must be tracked accurately for monthly/quarterly GST filing.',
         'GST Rate Dropdown (5%, 12%, 18%, 28%)'],
        ['Active Machines', "Count of machines with Status = 'Active'",
         'Operational machines generating revenue. Inactive machines need maintenance attention.',
         'None (static value)']
    ]
    
    for i, row_data in enumerate(kpi_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i+1].cells[j].text = cell_data
    
    doc.add_paragraph()
    
    # ===================== SECTION 3: CHARTS =====================
    doc.add_heading('3. Dashboard Charts', level=1)
    
    chart_table = doc.add_table(rows=5, cols=4)
    chart_table.style = 'Table Grid'
    
    chart_headers = ['Chart Type', 'Data Source', 'Calculation Method', 'Business Purpose']
    for i, header in enumerate(chart_headers):
        cell = chart_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'E2EFDA')
    
    chart_data = [
        ['Pie Chart (Category Distribution)', 'Stock + Products tables',
         'Σ (Stock × Cost) grouped by Category',
         'Shows inventory distribution - identify which product categories dominate your investment'],
        ['Area Chart (30-Day Trend)', 'Historical reconstruction from Sales + Refills',
         "Today's Value - Refills + Sales going backward day by day",
         'Visualize inventory value trend over time - detect patterns, seasonality, anomalies'],
        ['Bar Chart (Stock In/Out)', 'Refills + Sales logs',
         'Daily sum of Refill.Qty (In) vs Sales.Qty (Out)',
         'Track daily stock movement - identify busy days, plan staffing for refills'],
        ['Machine Comparison Bar', 'Stock per machine calculation',
         'Same as Total Stock Value formula, calculated per machine',
         'Compare machine performance - identify high/low performers for optimization']
    ]
    
    for i, row_data in enumerate(chart_data):
        for j, cell_data in enumerate(row_data):
            chart_table.rows[i+1].cells[j].text = cell_data
    
    doc.add_paragraph()
    
    # ===================== SECTION 4: INVENTORY PAGE =====================
    doc.add_heading('4. Inventory Page Components', level=1)
    
    doc.add_heading('Product Master Tab', level=2)
    p = doc.add_paragraph()
    p.add_run('Landed Cost Calculation: ').bold = True
    p.add_run('Landed_Cost = Unit_Cost × (1 + GST_Rate)')
    doc.add_paragraph('Example: If Unit_Cost = ₹100 and GST = 18%, then Landed_Cost = ₹100 × 1.18 = ₹118')
    doc.add_paragraph('This represents the true cost of acquiring the product including taxes.')
    
    doc.add_heading('Purchase Orders Tab', level=2)
    doc.add_paragraph('Displays all vendor purchases with status tracking (Pending, In Transit, Delivered).')
    doc.add_paragraph('Total Cost = Qty × PO_Price (from Excel)')
    
    doc.add_heading('KPIs on Inventory Page', level=2)
    inv_table = doc.add_table(rows=4, cols=3)
    inv_table.style = 'Table Grid'
    
    inv_headers = ['KPI', 'Formula', 'Meaning']
    for i, header in enumerate(inv_headers):
        cell = inv_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'FCE4D6')
    
    inv_data = [
        ['Total SKUs', 'products.length', 'Number of unique product types in catalog'],
        ['Pending POs', "purchases.filter(Status = 'Pending' OR 'In Transit').length", 'Orders waiting for delivery'],
        ['Total Value', 'Σ (Landed_Cost × Total_Stock)', 'Total inventory investment value']
    ]
    
    for i, row_data in enumerate(inv_data):
        for j, cell_data in enumerate(row_data):
            inv_table.rows[i+1].cells[j].text = cell_data
    
    doc.add_paragraph()
    
    # ===================== SECTION 5: MACHINES PAGE =====================
    doc.add_heading('5. Machines Page Components', level=1)
    
    doc.add_heading('Machine Cards', level=2)
    mach_table = doc.add_table(rows=4, cols=3)
    mach_table.style = 'Table Grid'
    
    mach_headers = ['Component', 'Calculation', 'Visual Indicator']
    for i, header in enumerate(mach_headers):
        cell = mach_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'DDEBF7')
    
    mach_data = [
        ['Stock Value', 'Σ (Current_Stock × Landed_Cost) for products in this machine', 'Displayed as ₹X,XXX'],
        ['Fill Level', '(Total_Items / 300) × 100%', 'Progress bar: Green (>70%), Yellow (30-70%), Red (<30%)'],
        ['Status Dot', 'Based on machine Status field', 'Green = Active, Yellow = Warning, Red = Critical']
    ]
    
    for i, row_data in enumerate(mach_data):
        for j, cell_data in enumerate(row_data):
            mach_table.rows[i+1].cells[j].text = cell_data
    
    doc.add_heading('Sell Button Action', level=2)
    doc.add_paragraph('When clicked, the "Sell 1" button:')
    doc.add_paragraph('1. Calls POST /api/sell with {machineId, productId, qty: 1, price: MRP}')
    doc.add_paragraph('2. Backend decrements Current_Stock by 1')
    doc.add_paragraph('3. Logs a new entry in Sales_Log with today\'s date')
    doc.add_paragraph('4. Uploads updated Excel to SharePoint')
    
    doc.add_paragraph()
    
    # ===================== SECTION 6: RESTOCK PAGE =====================
    doc.add_heading('6. Restock Page Components', level=1)
    
    doc.add_heading('Alert Classification Logic', level=2)
    alert_table = doc.add_table(rows=4, cols=3)
    alert_table.style = 'Table Grid'
    
    alert_headers = ['Status', 'Condition', 'Action Required']
    for i, header in enumerate(alert_headers):
        cell = alert_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'FFF2CC')
    
    alert_data = [
        ['Critical (Red)', 'Current_Stock < 10 units', 'Immediate restocking required - potential lost sales'],
        ['Low Stock (Yellow)', 'Current_Stock < Reorder_Level', 'Plan restocking within 2-3 days'],
        ['Safe (Green)', 'Current_Stock >= Reorder_Level', 'No action needed']
    ]
    
    for i, row_data in enumerate(alert_data):
        for j, cell_data in enumerate(row_data):
            alert_table.rows[i+1].cells[j].text = cell_data
    
    doc.add_heading('Refill Button Action', level=2)
    doc.add_paragraph('When clicked, the "Refill" button:')
    doc.add_paragraph('1. Calls POST /api/refill with {machineId, productId, qty: 50}')
    doc.add_paragraph('2. Backend adds 50 to Current_Stock')
    doc.add_paragraph('3. Logs entry in Machine_Refill_Log with today\'s date')
    doc.add_paragraph('4. Uploads updated Excel to SharePoint')
    
    doc.add_paragraph()
    
    # ===================== SECTION 7: IMPROVEMENTS =====================
    doc.add_heading('7. Recommended Improvements', level=1)
    
    doc.add_heading('Critical (Must Fix for Production)', level=2)
    crit_table = doc.add_table(rows=6, cols=3)
    crit_table.style = 'Table Grid'
    
    crit_headers = ['Issue', 'Current State', 'Recommended Fix']
    for i, header in enumerate(crit_headers):
        cell = crit_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'F8CBAD')
    
    crit_data = [
        ['No Authentication', 'Anyone can access the system', 'Add JWT/OAuth login with role-based access'],
        ['No Input Validation', 'Accepts any data without checks', 'Validate all inputs server-side before processing'],
        ['Hardcoded Capacity', 'Fill Level uses fixed 300 max', 'Make max capacity configurable per machine'],
        ['No Error Messages', 'Failures happen silently', 'Show user-friendly error notifications'],
        ['Negative Stock Possible', 'Can sell more than available', 'Add stock check before allowing sale']
    ]
    
    for i, row_data in enumerate(crit_data):
        for j, cell_data in enumerate(row_data):
            crit_table.rows[i+1].cells[j].text = cell_data
    
    doc.add_heading('High Priority Features', level=2)
    high_table = doc.add_table(rows=8, cols=2)
    high_table.style = 'Table Grid'
    
    high_headers = ['Feature', 'Current Status → Required']
    for i, header in enumerate(high_headers):
        cell = high_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'FFE699')
    
    high_data = [
        ['Search Functionality', 'UI exists but does not filter → Implement actual search logic'],
        ['Add Product Button', 'UI only → Connect to backend API for CRUD'],
        ['Create PO Button', 'UI only → Generate actual purchase orders'],
        ['Edit/Delete Buttons', 'UI only → Implement update/delete operations'],
        ['Export Report', 'UI only → Generate Excel/PDF exports'],
        ['Trend Percentages', 'Shows 0% always → Calculate actual week-over-week change'],
        ['Date Range Picker', 'Uses system date → Let user select custom date ranges']
    ]
    
    for i, row_data in enumerate(high_data):
        for j, cell_data in enumerate(row_data):
            high_table.rows[i+1].cells[j].text = cell_data
    
    doc.add_heading('Medium Priority Enhancements', level=2)
    med_list = [
        'Real-time WebSocket - Instant updates instead of 5-second polling',
        'Mobile Responsive Design - Better experience on phones/tablets',
        'Push Notifications - Alert when stock hits critical level',
        'Multi-user Support - Track which user performed each action',
        'Audit Log - Complete history of all changes for compliance',
        'Dashboard Customization - Let users arrange and resize widgets'
    ]
    for item in med_list:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Future Roadmap (Nice to Have)', level=2)
    future_list = [
        'Predictive Analytics - AI to forecast when stock will run out',
        'Auto-PO Generation - Automatically create PO when stock hits reorder level',
        'Payment Integration - Track vendor payments and reconciliation',
        'Multi-warehouse Support - Scale beyond vending machines to warehouses',
        'Barcode/QR Scanning - Quick product lookup via camera',
        'Customer Analytics - Track popular products and buying patterns'
    ]
    for item in future_list:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph()
    
    # ===================== SECTION 8: TECHNICAL DETAILS =====================
    doc.add_heading('8. Technical Implementation Details', level=1)
    
    doc.add_heading('Backend API Endpoints', level=2)
    api_table = doc.add_table(rows=4, cols=4)
    api_table.style = 'Table Grid'
    
    api_headers = ['Endpoint', 'Method', 'Purpose', 'Parameters']
    for i, header in enumerate(api_headers):
        cell = api_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D0CECE')
    
    api_data = [
        ['/api/dashboard', 'GET', 'Fetch all data for display', 'None'],
        ['/api/sell', 'POST', 'Record a sale transaction', 'machineId, productId, qty, price'],
        ['/api/refill', 'POST', 'Record a refill operation', 'machineId, productId, qty']
    ]
    
    for i, row_data in enumerate(api_data):
        for j, cell_data in enumerate(row_data):
            api_table.rows[i+1].cells[j].text = cell_data
    
    doc.add_heading('Excel Sheet Mapping', level=2)
    sheet_table = doc.add_table(rows=8, cols=3)
    sheet_table.style = 'Table Grid'
    
    sheet_headers = ['Internal Key', 'Excel Sheet Name', 'Purpose']
    for i, header in enumerate(sheet_headers):
        cell = sheet_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D0CECE')
    
    sheet_data = [
        ['Products', 'Product_Master', 'Product catalog with prices and GST'],
        ['Machines', 'Machine_Master', 'List of vending machines and locations'],
        ['Stock', 'Current_Stock', 'Current inventory levels per machine'],
        ['Sales', 'Sales_Log', 'Transaction history of all sales'],
        ['Purchases', 'Vendor_Purchase', 'Purchase order records'],
        ['Refills', 'Machine_Refill_Log', 'Restocking history'],
        ['Vendors', 'Vendor_Master', 'Supplier information']
    ]
    
    for i, row_data in enumerate(sheet_data):
        for j, cell_data in enumerate(row_data):
            sheet_table.rows[i+1].cells[j].text = cell_data
    
    doc.add_paragraph()
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph('--- End of Document ---')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save
    output_path = r'C:\Users\Bharani\OneDrive\Desktop\Tracker\VendBees_Documentation.docx'
    doc.save(output_path)
    print(f'Document saved to: {output_path}')
    return output_path

if __name__ == '__main__':
    create_document()
